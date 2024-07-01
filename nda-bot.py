import os
from docx import Document
from datetime import datetime
import smtplib
from email.message import EmailMessage

try:
    # Ask user whether it's for a developer or not
    is_developer = input("Is this NDA for a developer? (yes/no): ").lower().strip() == "yes"

    # Choose the appropriate NDA template
    if is_developer:
        docx_file_path = os.path.join('NDA', '_Developer NDA.docx')
        placeholders = {
            "[Developer's Name]": "Enter the developer's name: ",
            "[Software Developer/Web Developer]": "Enter the developer's position (e.g., Software Developer, "
                                                  "Web Developer): ",
            "[Developer's Email Address]": "Enter the developer's email address: ",
            "[Date]": ""

        }
    else:
        docx_file_path = os.path.join('NDA', '1_General NDA.docx')
        placeholders = {
            "[Employee Name]": "Enter the employee's name: ",
            "[Employee Email]": "Enter the employee's email: ",
            "[Employee Phone]": "Enter the employee's phone: ",
            "[Assistant/Website Security Manager/SEO Strategy Manager/Talent Acquisition Specialist/Figma Designer]":
                "Enter the employee's position (e.g., Assistant/Website Security Manager/SEO Strategy Manager/Talent "
                " Acquisition) : ",
            "[Date]": ""
        }

    # Read the selected NDA template
    doc = Document(docx_file_path)

    # Prompt the user for input for each placeholder
    user_inputs = {}
    for placeholder, prompt in placeholders.items():
        user_input = input(prompt)
        if placeholder == "[Date]":
            user_input = datetime.now().strftime('%B %d, %Y')  # Format the date as desired
        user_inputs[placeholder] = user_input
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, user_input)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, user_input)
    # Create a directory for updated NDA documents if it doesn't exist
    output_directory = 'updated-NDA'
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    # Construct the filename for the updated NDA document
    if is_developer:
        employee_name = user_inputs.get("[Developer's Name]", "Unknown")
    else:
        employee_name = user_inputs.get("[Employee Name]", "Unknown")
    file_name = f"{employee_name}-{datetime.now().date()}.docx"
    updated_docx_file_path = os.path.join(output_directory, file_name)

    # Save the modified document with updated placeholders
    doc.save(updated_docx_file_path)

    print("Updated NDA document created successfully:", updated_docx_file_path)
    email = EmailMessage()
    email['From'] = 'mughira3107@gmail.com'  # Your Gmail address
    employee_email = user_inputs.get("[Employee Email]", "")
    email['To'] = employee_email
    email['Cc'] = 'hafizmughira153@gmail.com'
    email['Subject'] = 'NDA Agreement'

    # Email body
    email.set_content('Signed NDA attached .')

    # Attach the updated NDA document
    with open(updated_docx_file_path, 'rb') as file:
        content = file.read()
        email.add_attachment(content, maintype='application', subtype='docx', filename=file_name)

    # Send the email using Gmail SMTP server
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login('mughira3107@gmail.com', 'nulp uegi fzpz sdvv')  # Your Gmail address and password
        smtp.send_message(email)

    print("Email sent successfully to:", employee_email)


except Exception as e:
    print(f"An error occurred: {str(e)}")
